import asyncio
import os
import time
import wave
from pathlib import Path

import numpy as np
import pyaudio
import requests
import torch
from audio_recorder_streamlit import audio_recorder
from docx import Document
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline

import streamlit as st

# Set the page config
st.set_page_config(
    page_title="Audio to Medical Letter Generator",
    layout="centered",
    initial_sidebar_state="auto",
)


# Check if GPU is available
device = "cuda" if torch.cuda.is_available() else "cpu"


# Load the Whisper ASR model
@st.cache_resource
def load_asr_model():
    return pipeline(
        # "automatic-speech-recognition",
        # "openai/whisper-medium",
        "automatic-speech-recognition",
        "openai/whisper-tiny",
        chunk_length_s=30,
        stride_length_s=5,
        return_timestamps=False,
        device=device,
    )


asr_pipe = load_asr_model()


# Load the medical letter generator model
@st.cache_resource
def generate_medical_letter(letter_type, patient_summary):
    # Prepare prompt with structured instructions
    prompt = f"""
You are a medical assistant specialized in generating professional medical {letter_type}.
Generate a professional medical letter based on the following structured format using the text provided.

Structured Format:

Patient Name: 
Date of Birth: 
Medical Record Number: 

Date of Visit: 
Primary Complaint: 
Referring Physician: 

Details:
- Patient's Medical History:
- Examination Findings:
- Diagnosis:
- Treatment Plan:
- Follow-Up Instructions:

Physician's Name: 
Signature:

Medical Report Text:


Patient Summary:
{patient_summary}
"""

    with st.spinner("Generating your medical letter..."):
        try:
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={"model": "mistral", "prompt": prompt, "stream": False},
                timeout=300,
            )
            response.raise_for_status()
            result = response.json()
            return result.get("response", "Error: No response from TinyLlama.")
        except Exception as e:
            return f"Error generating medical letter: {str(e)}"


# Save the letter as a Word document
def save_letter_to_word(letter_content, file_name):
    doc = Document()
    doc.add_paragraph(letter_content)
    temp_file_path = os.path.join("temp_dir", file_name)
    os.makedirs("temp_dir", exist_ok=True)
    doc.save(temp_file_path)

    return temp_file_path


# Handle transcription
def handle_file_transcription(uploaded_file):
    with st.spinner("Transcribing..."):
        start_time = time.time()
        temp_dir = "temp_dir"
        os.makedirs(temp_dir, exist_ok=True)
        temp_file_path = os.path.join(temp_dir, uploaded_file.name)

        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        transcription = asr_pipe(temp_file_path)
        end_time = time.time()
        st.success(
            f"Transcription completed in {round(end_time - start_time, 2)} seconds!"
        )

        return transcription["text"]


# Ensure session state keys are initialized
if "transcription_text" not in st.session_state:
    st.session_state["transcription_text"] = ""
if "generated_letter" not in st.session_state:
    st.session_state["generated_letter"] = ""


# Handle recorded audio transcription
def handle_recorded_transcription_1(temp_file_path):
    with st.spinner("Transcribing..."):
        start_time = time.time()
        try:
            transcription = asr_pipe(temp_file_path)
            st.success("Transcription completed!")

            # Store transcription in session state
            st.session_state["transcription_text"] = transcription["text"]
            end_time = time.time()
            st.write(f"Time taken: {round(end_time - start_time, 2)} seconds")

            # Clean up temporary file
            os.remove(temp_file_path)

            # Return the transcription text
            return transcription["text"]
        except Exception as e:
            st.error(f"An error occurred during transcription: {e}")
            return None


def save_transcription_to_word(transcription_content, file_name):
    try:
        # Create a new Word document
        doc = Document()
        doc.add_paragraph(transcription_content)

        # Save the document to the current directory
        temp_file_path = os.path.join(os.getcwd(), file_name)
        doc.save(temp_file_path)

        # Debugging output
        st.write(f"File saved at: {temp_file_path}")
        return temp_file_path
    except Exception as e:
        st.error(f"Error saving transcription file: {e}")
        return None


# Define available medical letter types
medical_letters = [
    "Referral Letter",
    "Discharge Summary",
    "Sick Leave Certificate",
    "Specialist Report",
    "Medical Fitness Certificate",
    "Operation Report",
    "Insurance Claim Letter",
    "Follow-up Letter",
    "Medical History Summary",
]


# Define the app layout
def main():
    st.title("Audio to Medical Letter Generator")

    # Tabs for functionalities
    tabs = st.tabs(
        ["Upload Audio", "Record Audio", "Generate Letter", "Real Time "]
    )

    # Tab for uploading audio and transcription
    with tabs[0]:
        st.subheader("Upload and Transcribe Audio")
        uploaded_file = st.file_uploader(
            "Upload an audio file", type=["mp3", "wav", "ogg", "m4a"]
        )

        if uploaded_file:
            st.audio(uploaded_file)
            if st.button(
                "Transcribe Uploaded File", key="transcribe_button_upload"
            ):
                transcription_text = handle_file_transcription(uploaded_file)
                st.session_state["transcription_output"] = transcription_text

        # Display the transcription output if available
        transcription_text = st.session_state.get("transcription_output", "")
        st.text_area(
            "Transcription Output",
            value=transcription_text,
            height=300,
            key="transcription_output",
        )

        # Allow downloading transcription as a Word document
        if transcription_text.strip():
            transcription_doc_path = save_letter_to_word(
                transcription_text, "transcription.docx"
            )
            with open(transcription_doc_path, "rb") as f:
                st.download_button(
                    "Download Transcription as Word Document",
                    data=f,
                    file_name="transcription.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_transcription_word",
                )

        # Medical letter generation section
        st.subheader("Generate Medical Letter from Transcription")
        selected_letter = st.selectbox(
            "Select the type of medical letter to generate:",
            medical_letters,
            index=0,
            key="select_medical_letter",
        )
        st.write("You selected:", selected_letter)
        patient_summary = st.text_area(
            "Edit Transcription or Add Details",
            value=transcription_text,
            height=300,
            key="editable_transcription",
        )

        if st.button(
            "Generate Medical Letter",
            key="generate_medical_letter_button_tab_0",
        ):
            if not patient_summary.strip():

                st.error("Please provide patient details and summary.")
            else:

                letter_content = generate_medical_letter(
                    selected_letter, patient_summary
                )

                st.session_state["generated_letter"] = letter_content

        # Display the generated medical letter if available
        generated_letter = st.session_state.get("generated_letter", "")

        if generated_letter:
            st.subheader("Generated Medical Letter")
            st.text_area(
                "Your Medical Letter:",
                value=generated_letter,
                height=400,
                key="generated_letter_tab_0",
            )

            # Save and allow download of the generated medical letter
            letter_file_name = (
                f"{selected_letter.replace(' ', '_').lower()}_letter.docx"
            )
            temp_file_path = save_letter_to_word(
                generated_letter, letter_file_name
            )
            with open(temp_file_path, "rb") as f:
                st.download_button(
                    "Download Medical Letter as Word Document",
                    data=f,
                    file_name=letter_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_generated_letter_tab_0",
                )

    # Record audio tab

    with tabs[1]:
        st.header("🎙️ Record Audio")
        # st.markdown("Click the mic icon below to start or stop recording.")

        audio_bytes = audio_recorder(
            text="Click the mic icon to start or stop the recording",
            recording_color="#e8b62c",
            icon_name="microphone",
            icon_size="3x",
        )

        if audio_bytes:
            st.audio(audio_bytes, format="audio/wav")

            temp_dir = "temp_dir"
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, "recorded_audio.wav")
            with open(temp_file_path, "wb") as f:
                f.write(audio_bytes)

            if st.button("📝 Transcribe Recorded Audio"):
                handle_recorded_transcription_1(temp_file_path)

        # Display transcription text if available
        if st.session_state.get("transcription_text"):
            st.subheader("🗒️ Transcription")
            transcription_text = st.session_state["transcription_text"]
            st.text_area(
                "Review or edit the transcription below:",
                value=transcription_text,
                height=300,
            )
            transcription_doc_path = save_letter_to_word(
                transcription_text, "transcription.docx"
            )
            with open(transcription_doc_path, "rb") as f:
                st.download_button(
                    "⬇️ Download Transcription as Word Document",
                    data=f,
                    file_name="transcription.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            st.subheader("✍️ Generate Medical Letter")
            selected_letter = st.selectbox(
                "Select the type of medical letter to generate:",
                medical_letters,
                index=0,
            )
            st.write(f"You selected: **{selected_letter}**")
            patient_summary = st.text_area(
                "Edit transcription or add patient details:",
                value=transcription_text,
                height=300,
            )
            if st.button("📄 Generate Medical Letter"):
                if not patient_summary.strip():
                    st.error("❌ Please provide patient details and summary.")
                else:
                    letter_content = generate_medical_letter(
                        selected_letter, patient_summary
                    )
                    st.session_state["generated_letter"] = letter_content

        if st.session_state.get("generated_letter"):
            st.subheader("📄 Generated Medical Letter")
            generated_letter = st.session_state["generated_letter"]

            st.text_area(
                "Your Medical Letter:",
                value=generated_letter,
                height=400,
            )
            letter_file_name = (
                f"{selected_letter.replace(' ', '_').lower()}_letter.docx"
            )
            temp_file_path = save_letter_to_word(
                generated_letter, letter_file_name
            )
            with open(temp_file_path, "rb") as f:
                st.download_button(
                    "⬇️ Download Medical Letter as Word Document",
                    data=f,
                    file_name=letter_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    # Generate Letter tab
    with tabs[2]:
        st.subheader("Generate Medical Letter")
        transcription_text = st.session_state.get("transcription_output", "")
        patient_summary = st.text_area(
            "Patient Summary",
            value=transcription_text,
            height=300,
            key="patient_summary_2",
        )
        letter_type = st.selectbox(
            "Select the type of medical letter:",
            medical_letters,
            key="letter_type_selector_2",
        )

        if st.button(
            "Generate Medical Letter", key="generate_medical_letter_button_2"
        ):
            if not patient_summary.strip():
                st.error("Please provide patient details and summary.")
            else:
                letter_content = generate_medical_letter(
                    letter_type, patient_summary
                )
                st.subheader("Generated Medical Letter")
                st.text_area(
                    "Your Medical Letter:",
                    value=letter_content,
                    height=400,
                    key="generated_letter_2",
                )
                # Save and allow download of the generated medical letter
                letter_file_name = (
                    f"{letter_type.replace(' ', '_').lower()}_letter.docx"
                )
                temp_file_path = save_letter_to_word(
                    letter_content, letter_file_name
                )

                with open(temp_file_path, "rb") as f:
                    st.download_button(
                        "Download Medical Letter as Word Document",
                        data=f,
                        file_name=letter_file_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_generated_letter_2",
                    )

    with tabs[3]:
        # UI HEADER & SIDEBAR
        st.header("⏱️ Real-Time Dictation")

        st.sidebar.subheader("Real-Time Settings")
        RATE = int(st.sidebar.text_input("Sample Rate", value="16000"))
        FRAMES_PER_BUFFER = int(
            st.sidebar.text_input("Buffer Size", value="3200")
        )

        # SESSION STATE SETUP
        if "run_rt" not in st.session_state:
            st.session_state["run_rt"] = False
        if "real_time_text" not in st.session_state:
            st.session_state["real_time_text"] = ""
        if "generated_letter_rt" not in st.session_state:
            st.session_state["generated_letter_rt"] = ""

        # TOGGLE FUNCTION
        def toggle_rt():
            st.session_state["run_rt"] = not st.session_state["run_rt"]

        # BUTTONS
        if not st.session_state["run_rt"]:
            st.button("▶️ Start Real-Time Dictation", on_click=toggle_rt)
        else:
            st.button("⏹ Stop Dictation", on_click=toggle_rt)

        # REAL-TIME TRANSCRIPTION LOOP
        if st.session_state["run_rt"]:
            p = pyaudio.PyAudio()
            stream = p.open(
                format=pyaudio.paInt16,
                channels=1,
                rate=RATE,
                input=True,
                frames_per_buffer=FRAMES_PER_BUFFER,
            )

            frames = []
            status_placeholder = st.empty()
            live_text_placeholder = st.empty()

            while st.session_state["run_rt"]:
                try:
                    data = stream.read(
                        FRAMES_PER_BUFFER, exception_on_overflow=False
                    )
                    frames.append(data)

                    # Process every 10 seconds
                    if len(frames) * FRAMES_PER_BUFFER / RATE > 10:
                        audio_data = b"".join(frames)
                        frames = []  # Reset buffer

                        # Save to temp file
                        with wave.open("temp_rt_audio.wav", "wb") as wf:
                            wf.setnchannels(1)
                            wf.setsampwidth(p.get_sample_size(pyaudio.paInt16))
                            wf.setframerate(RATE)
                            wf.writeframes(audio_data)

                        # Convert to numpy
                        with wave.open("temp_rt_audio.wav", "rb") as wf:
                            raw_audio = wf.readframes(wf.getnframes())
                            audio_array = (
                                np.frombuffer(raw_audio, dtype=np.int16).astype(
                                    np.float32
                                )
                                / 32768.0
                            )

                        # Transcribe
                        result = asr_pipe(audio_array)
                        st.session_state["real_time_text"] += (
                            result["text"] + " "
                        )
                        os.remove("temp_rt_audio.wav")

                        # Live UI update
                        status_placeholder.info("Listening...")
                        live_text_placeholder.text_area(
                            "Live Transcription",
                            st.session_state["real_time_text"],
                            height=300,
                        )

                except Exception as e:
                    status_placeholder.error(f"Error during transcription: {e}")
                    break

            stream.stop_stream()
            stream.close()
            p.terminate()

        else:
            # Display transcription text if not running
            st.text_area(
                "Live Transcription",
                st.session_state["real_time_text"],
                height=300,
            )

        # DOWNLOAD TRANSCRIPTION
        if st.session_state["real_time_text"].strip():
            transcription_file = save_letter_to_word(
                st.session_state["real_time_text"],
                "real_time_transcription.docx",
            )
            with open(transcription_file, "rb") as f:
                st.download_button(
                    "💾 Download Transcription",
                    f,
                    file_name="real_time_transcription.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

        # GENERATE LETTER SECTION
        st.subheader("Generate Medical Letter from Real-Time Transcription")
        selected_letter_rt = st.selectbox(
            "Select Letter Type", medical_letters, key="rt_letter_type"
        )
        patient_summary_rt = st.text_area(
            "Edit Transcription or Add Details",
            value=st.session_state["real_time_text"],
            height=300,
        )

        if st.button("Generate Medical Letter", key="generate_letter_rt"):
            if not patient_summary_rt.strip():
                st.error("Please provide patient details and summary.")
            else:
                letter_content_rt = generate_medical_letter(
                    selected_letter_rt, patient_summary_rt
                )
                st.session_state["generated_letter_rt"] = letter_content_rt

        if st.session_state["generated_letter_rt"]:
            st.text_area(
                "Generated Medical Letter",
                st.session_state["generated_letter_rt"],
                height=400,
            )
            letter_file_name_rt = (
                f"{selected_letter_rt.replace(' ', '_').lower()}_letter.docx"
            )
            temp_file_path_rt = save_letter_to_word(
                st.session_state["generated_letter_rt"], letter_file_name_rt
            )
            with open(temp_file_path_rt, "rb") as f:
                st.download_button(
                    "📄 Download Medical Letter",
                    f,
                    file_name=letter_file_name_rt,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )


if __name__ == "__main__":

    main()